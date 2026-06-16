import os
import re
from datetime import datetime, timedelta
import docx
import openpyxl
from openpyxl.styles import PatternFill, Font, Alignment, Border, Side
from openpyxl.utils import get_column_letter
from holidays import get_holidays_by_year

# -------------------------------------------------------------------------
# Türkçe ay adları → ay numarası eşleştirmesi
# -------------------------------------------------------------------------
# Normalize edilmiş (ASCII benzeri) ay adları eşleştirmesi
# normalize_turkish() fonksiyonundan geçtikten sonra bu değerlerle eşleşir
TURKISH_MONTHS = {
    # normalize_turkish() sonucu → ay numarası
    "OCAK": 1,
    "SUBAT": 2,     # ŞUBAT → SUBAT
    "MART": 3,
    "NISAN": 4,     # NİSAN → NISAN
    "MAYIS": 5,
    "HAZIRAN": 6,
    "TEMMUZ": 7,
    "AGUSTOS": 8,   # AĞUSTOS → AGUSTOS
    "EYLUL": 9,     # EYLÜL → EYLUL
    "EKIM": 10,     # EKİM → EKIM
    "KASIM": 11,
    "ARALIK": 12,
}

def normalize_turkish(text: str) -> str:
    """
    Türkçe harfleri ASCII karşılıklarına dönüştürür VE büyük harfe çevirir.
    Sıralama önemli: önce küçük+büyük Türkçe karakterleri değiştir, sonra upper().
    """
    if not text:
        return ""
    # Karakter çiftleri: Türkçe → ASCII (küçük ve büyük)
    mappings = [
        ("ı", "i"), ("İ", "I"),
        ("ğ", "g"), ("Ğ", "G"),
        ("ş", "s"), ("Ş", "S"),
        ("ç", "c"), ("Ç", "C"),
        ("ö", "o"), ("Ö", "O"),
        ("ü", "u"), ("Ü", "U"),
    ]
    result = text
    for tr_char, ascii_char in mappings:
        result = result.replace(tr_char, ascii_char)
    return result.upper()


def parse_week_cell(cell_value: str, academic_year: str) -> tuple[datetime, datetime] | None:
    """
    '14-18 EYLÜL', '30 KASIM - 4 ARALIK', '25-29 OCAK' gibi
    hafta aralığı metinlerini (start_date, end_date) çiftine çevirir.
    
    Dönem başlangıç yılı (Eylül-Aralık) ve bitiş yılı (Ocak-Haziran) ayrı hesaplanır.
    """
    if not cell_value:
        return None
    
    text = str(cell_value).strip().upper()
    # Türkçe normalize
    text_norm = normalize_turkish(text)
    
    # Yılları ayır: 'XXXX-YYYY' → start_year, end_year
    try:
        parts = academic_year.split("-")
        start_year = int(parts[0])  # 2026
        end_year = int(parts[1])    # 2027
    except Exception:
        start_year = datetime.now().year
        end_year = start_year + 1

    # --- PATTERN 1: "14-18 EYLÜL" (tek ay) ---
    m = re.match(r"(\d{1,2})\s*[-–]\s*(\d{1,2})\s+([A-ZÇĞIİÖŞÜa-zçğışöü]+)", text)
    if m:
        day1 = int(m.group(1))
        day2 = int(m.group(2))
        month_str = normalize_turkish(m.group(3))
        month_num = TURKISH_MONTHS.get(month_str)
        if month_num:
            year = end_year if month_num <= 6 else start_year
            try:
                return datetime(year, month_num, day1), datetime(year, month_num, day2)
            except ValueError:
                return None

    # --- PATTERN 2: "30 KASIM - 4 ARALIK" veya "28 OCAK - 1 ŞUBAT" (farklı aylar) ---
    m = re.match(
        r"(\d{1,2})\s+([A-ZÇĞIİÖŞÜa-zçğışöü]+)\s*[-–]\s*(\d{1,2})\s+([A-ZÇĞIİÖŞÜa-zçğışöü]+)",
        text,
    )
    if m:
        day1 = int(m.group(1))
        month_str1 = normalize_turkish(m.group(2))
        day2 = int(m.group(3))
        month_str2 = normalize_turkish(m.group(4))
        month1 = TURKISH_MONTHS.get(month_str1)
        month2 = TURKISH_MONTHS.get(month_str2)
        if month1 and month2:
            year1 = end_year if month1 <= 6 else start_year
            year2 = end_year if month2 <= 6 else start_year
            try:
                return datetime(year1, month1, day1), datetime(year2, month2, day2)
            except ValueError:
                return None

    return None


def week_overlaps_holiday(week_start: datetime, week_end: datetime, holiday: dict) -> bool:
    """
    Hafta aralığı [week_start, week_end] ile tatil aralığı [holiday_start, holiday_end]
    kesişiyor mu kontrol et.
    """
    try:
        h_start = datetime.strptime(holiday["start"], "%Y-%m-%d")
        h_end = datetime.strptime(holiday["end"], "%Y-%m-%d")
    except Exception:
        return False
    # Klasik aralık kesişim kontrolü
    return week_start <= h_end and week_end >= h_start


def find_overlapping_holidays(week_start: datetime, week_end: datetime, holidays: list) -> list:
    """Bu haftayla çakışan tatillerin listesini döndür."""
    return [h for h in holidays if week_overlaps_holiday(week_start, week_end, h)]


# -------------------------------------------------------------------------
# Excel İşleyici
# -------------------------------------------------------------------------
HOLIDAY_FILL   = PatternFill(start_color="FFFF6B6B", end_color="FFFF6B6B", fill_type="solid")  # kırmızı
HOLIDAY_FONT   = Font(bold=True, color="FFFFFFFF")
PARTIAL_FILL   = PatternFill(start_color="FFFFD93D", end_color="FFFFD93D", fill_type="solid")   # sarı
PARTIAL_FONT   = Font(bold=True, color="FF333333")


def _find_week_column(ws) -> int | None:
    """
    Hafta tarih formatını ('14-18 EYLÜL', '30 KASIM - 4 ARALIK' vb.) içeren
    sütun indeksini (1 tabanlı) otomatik tespit eder.
    İlk 20 satırı tarar ve en çok eşleşme olan sütunu döndürür.
    """
    import re
    col_hits = {}  # col_index -> hit_count
    for row in ws.iter_rows(max_row=20):
        for cell in row:
            if cell.value is None:
                continue
            text = str(cell.value).strip()
            # Hafta formatı: "14-18 EYLÜL" veya "30 KASIM - 4 ARALIK"
            if re.search(r'\d{1,2}\s*[-–]\s*\d{1,2}', text):
                col_hits[cell.column] = col_hits.get(cell.column, 0) + 1
    if not col_hits:
        return None
    # En çok eşleşme olan sütunu seç
    return max(col_hits, key=col_hits.get)


def process_excel_plan(excel_path: str, holidays: list, output_path: str) -> str:
    wb = openpyxl.load_workbook(excel_path)
    ws = wb.active

    max_col = ws.max_column

    # Hafta tarihlerinin bulunduğu sütunu otomatik tespit et
    week_col = _find_week_column(ws)
    if week_col is None:
        week_col = 1  # Bulunamazsa ilk sütunu dene

    holiday_rows = []   # (row_index, [holiday_names], is_full_holiday)

    for row in ws.iter_rows():
        # Hafta sütununu oku (otomatik tespit edilmiş)
        if week_col - 1 >= len(row):
            continue
        week_cell = row[week_col - 1]
        cell_val = week_cell.value
        if cell_val is None:
            continue

        parsed = parse_week_cell(str(cell_val), _get_academic_year(holidays))
        if parsed is None:
            continue

        week_start, week_end = parsed
        overlapping = find_overlapping_holidays(week_start, week_end, holidays)
        if not overlapping:
            continue

        # Tatil haftasının tamamı mı, yoksa kısmen mi?
        h_start = min(datetime.strptime(h["start"], "%Y-%m-%d") for h in overlapping)
        h_end   = max(datetime.strptime(h["end"],   "%Y-%m-%d") for h in overlapping)
        is_full = h_start <= week_start and h_end >= week_end

        holiday_rows.append((week_cell.row, overlapping, is_full))

    # Renklendirme
    for (row_idx, overlapping, is_full) in holiday_rows:
        fill = HOLIDAY_FILL if is_full else PARTIAL_FILL
        font = HOLIDAY_FONT if is_full else PARTIAL_FONT

        for col_idx in range(1, max_col + 1):
            cell = ws.cell(row=row_idx, column=col_idx)
            cell.fill = fill
            # Varsa metin içeriklerini koruyalım, yalnızca font güncelle
            if cell.value is not None:
                cell.font = font

        # Tatil adını son sütunun yanına ekle
        note_col = max_col + 1
        names = " / ".join(h["name"] for h in overlapping)
        note_cell = ws.cell(row=row_idx, column=note_col, value=f"🎉 {names}")
        note_cell.fill = fill
        note_cell.font = font
        note_cell.alignment = Alignment(horizontal="left", vertical="center", wrap_text=True)

    wb.save(output_path)
    return output_path


def _get_academic_year(holidays: list) -> str:
    """Tatil listesinden akademik yılı tahmin et."""
    for h in holidays:
        try:
            year = datetime.strptime(h["start"], "%Y-%m-%d").year
            return f"{year-1}-{year}" if datetime.strptime(h["start"], "%Y-%m-%d").month <= 6 else f"{year}-{year+1}"
        except Exception:
            pass
    return f"{datetime.now().year}-{datetime.now().year + 1}"


# -------------------------------------------------------------------------
# Word İşleyici
# -------------------------------------------------------------------------
from docx.shared import RGBColor, Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy


def _set_cell_bg(cell, hex_color: str):
    """Word tablo hücresine arka plan rengi uygula."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    # Varsa eski shd'yi sil
    for existing in tcPr.findall(qn("w:shd")):
        tcPr.remove(existing)
    tcPr.append(shd)


def process_word_plan(docx_path: str, holidays: list, output_path: str) -> str:
    doc = docx.Document(docx_path)
    academic_year = _get_academic_year(holidays)

    for table in doc.tables:
        for row in table.rows:
            if not row.cells:
                continue
            first_cell_text = row.cells[0].text.strip()
            parsed = parse_week_cell(first_cell_text, academic_year)
            if parsed is None:
                continue

            week_start, week_end = parsed
            overlapping = find_overlapping_holidays(week_start, week_end, holidays)
            if not overlapping:
                continue

            h_start = min(datetime.strptime(h["start"], "%Y-%m-%d") for h in overlapping)
            h_end   = max(datetime.strptime(h["end"],   "%Y-%m-%d") for h in overlapping)
            is_full = h_start <= week_start and h_end >= week_end

            bg_color = "FF6B6B" if is_full else "FFD93D"   # kırmızı / sarı

            names = " / ".join(h["name"] for h in overlapping)

            for cell in row.cells:
                _set_cell_bg(cell, bg_color)
                # Metni kalın yap
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.bold = True
                        if is_full:
                            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

            # Son hücreye tatil adını ekle
            last_cell = row.cells[-1]
            existing = last_cell.text.strip()
            if existing:
                last_cell.paragraphs[0].add_run(f"  🎉 {names}")
            else:
                last_cell.paragraphs[0].text = f"🎉 {names}"

    doc.save(output_path)
    return output_path


# -------------------------------------------------------------------------
# Ana Giriş Noktası
# -------------------------------------------------------------------------
def process_plan_file(plan_path: str, academic_year: str) -> str:
    """Yıllık planı tatillerle eşleştirerek işle ve çıktı dosyasını döndür."""
    holidays = get_holidays_by_year(academic_year)

    filename = os.path.basename(plan_path)
    output_filename = "duzenlenmis_" + filename
    output_dir = os.path.normpath(os.path.join(os.path.dirname(plan_path), "..", "processed"))
    os.makedirs(output_dir, exist_ok=True)
    output_path = os.path.join(output_dir, output_filename)

    if plan_path.lower().endswith(".xlsx"):
        process_excel_plan(plan_path, holidays, output_path)
    elif plan_path.lower().endswith(".docx"):
        process_word_plan(plan_path, holidays, output_path)
    else:
        # Desteklenmeyen format — olduğu gibi kopyala
        import shutil
        shutil.copy2(plan_path, output_path)

    return output_path
